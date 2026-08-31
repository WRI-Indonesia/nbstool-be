# application/utils/document_generator/v3/convert.py
#
# Turn the doc team's bracket-tagged draft into a docxtpl (jinja) template.
#
# The doc team authors and iterates the template in Word using human-readable placeholders:
#     [Site Characterisation: elevation class]   [User input: Project Proponent]
# Jinja tags do not survive Word editing (run splitting, smart quotes), so the bracket file stays
# the AUTHORING format and this converter is re-run per template revision:
#
#     python -m application.utils.document_generator.v3.convert \
#         "assets/Feasibility Study_Template Draft.docx" assets/feasibility_v3_template.docx
#
# What it does:
#   1. Every `[...]` placeholder becomes `{{ t("...") }}` -- one lookup function, no per-tag
#      variable names. The context's `t()` returns the mapped value or the literal bracket text,
#      so an unmapped or unfilled tag stays visible in the output for manual fill.
#   2. The land-cover "rows repeat for each remaining class" row and the species annex row become
#      `{%tr for %}` loops (see LOOP_ROWS).
#   3. `[only if X] ... [/only if X]` spans become `{% if %} ... {% endif %}`.
#
# Conversion is RUN-AWARE: only the runs a bracket actually spans are merged, so colored or
# otherwise formatted text around a placeholder keeps its own runs and formatting, and the
# replacement inherits the placeholder run's formatting (the first cut collapsed each bracketed
# paragraph into its first run, which flattened every colour in it to run 1's -- doc team
# noticed).
#
# The draft's own "Data Tag Reference" table (an authoring aid listing every tag) is DROPPED
# from the generated template (team decision 2026-08-31), together with its heading paragraph.
# It is removed BEFORE conversion, so occurrence numbering never sees its rows.

from __future__ import annotations

import re
import sys
from copy import deepcopy

from docx import Document

BRACKET = re.compile(r"\[([^\[\]]+)\]")
ONLY_IF = re.compile(r"\[only if ([^\[\]]+?)\]")
END_ONLY_IF = re.compile(r"\[/only if [^\[\]]*\]")

# Conditions used by [only if ...] spans -> context flag names.
CONDITIONS = {
    "keystone species present": "keystone_present",
}

# Repeating table rows. A row whose text contains `marker` is turned into a loop over `var`:
# a `{%tr for %}` row above, `{%tr endfor %}` below, and each `[tag]` cell replaced by the
# matching jinja expression. Sibling paragraphs saying "(rows repeat...)" are dropped.
LOOP_ROWS = [
    {
        "marker": "n land cover class",
        "var": "land_cover_rest",
        "cells": {
            "n land cover class": "{{ r.name }}",
            "n land cover class area (ha)": "{{ r.area }}",
            "n land cover class area (% of AOI)": "{{ r.pct }}",
        },
        # The "No" column of the repeated row counts on from the three fixed rows.
        "literal": {"..": "{{ loop.index + 3 }}"},
    },
    {
        "marker": "Species taxonomic class",
        "var": "species_rows",
        "cells": {
            "Species taxonomic class": "{{ r.taxon_class }}",
            "Scientific name": "{{ r.scientific_name }}",
            "Species IUCN Red List Category": "{{ r.redlist_category }}",
        },
        "literal": {"[n]": "{{ loop.index }}"},
    },
]

REPEAT_NOTE = "(rows repeat"


_WS = re.compile(r"\s+")


def _norm(tag: str) -> str:
    return _WS.sub(" ", tag).strip()


def _to_jinja(text: str, counters: dict | None = None, occurrences: dict | None = None) -> str:
    """Bracket placeholders in one text blob -> jinja, conditionals first.

    `occurrences` maps a normalised tag to `fn(n) -> jinja expression or None`, where `n` is the
    tag's 1-based occurrence number in DOCUMENT ORDER (tracked in `counters`, shared across the
    whole document). The Monitoring Plan repeats the same tag text once per ecosystem block, so
    only the position can say which ecosystem a given occurrence belongs to. None falls back to
    the generic `t()` lookup.
    """

    def _cond(match):
        flag = CONDITIONS.get(match.group(1).strip())
        # An unknown condition stays literal, visible in the output rather than silently eaten.
        return "{%% if %s %%}" % flag if flag else match.group(0)

    def _tag(match):
        tag = match.group(1)
        key = _norm(tag)
        if occurrences is not None and key in occurrences:
            counters[key] = counters.get(key, 0) + 1
            expr = occurrences[key](counters[key])
            if expr is not None:
                return "{{ %s }}" % expr
        # Escape quotes inside the tag text so the jinja string literal survives.
        return '{{ t("%s") }}' % tag.replace('"', '\\"')

    text = ONLY_IF.sub(_cond, text)
    text = END_ONLY_IF.sub("{% endif %}", text)
    return BRACKET.sub(_tag, text)


def _rewrite_paragraph(paragraph, new_text: str) -> None:
    """Put `new_text` into the paragraph's first run and drop the rest.

    Used where the WHOLE cell content is replaced (loop rows, injected control rows), so there
    is no surrounding formatting to keep.
    """
    if not paragraph.runs:
        if new_text:
            paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)


# One bracket token, [only if ...] and [/only if ...] included -- they are brackets too.
_TOKEN = re.compile(r"\[[^\[\]]+\]")


def _convert_paragraph(paragraph, counters=None, occurrences=None) -> None:
    """Convert one paragraph's brackets RUN-AWARE.

    Word splits a placeholder across runs at will, so each bracket's span of runs is merged into
    the run where it starts -- which guarantees the jinja tag is never split AND that the
    replacement keeps the placeholder run's own formatting. Runs outside a bracket are left
    untouched, so surrounding coloured or styled text survives conversion. Left-to-right order
    keeps the occurrence counters in document order.
    """
    scan_from = 0
    while True:
        runs = paragraph.runs
        text = "".join(run.text or "" for run in runs)
        match = _TOKEN.search(text, scan_from)
        if not match:
            return
        replacement = _to_jinja(match.group(0), counters, occurrences)
        start, end = match.span()
        if replacement == match.group(0):
            # An unknown [only if ...] condition stays literal by design; step past it.
            scan_from = end
            continue
        scan_from = start + len(replacement)

        bounds = []
        position = 0
        for run in runs:
            length = len(run.text or "")
            bounds.append((position, position + length))
            position += length
        first = next(i for i, (a, b) in enumerate(bounds) if a <= start < b)
        last = next(i for i, (a, b) in enumerate(bounds) if a < end <= b)

        prefix = (runs[first].text or "")[:start - bounds[first][0]]
        suffix = (runs[last].text or "")[end - bounds[last][0]:]
        if first == last:
            runs[first].text = prefix + replacement + suffix
        else:
            runs[first].text = prefix + replacement
            runs[last].text = suffix
            for run in runs[first + 1:last]:
                run.text = ""


def _convert_paragraphs(paragraphs, counters=None, occurrences=None) -> None:
    for paragraph in paragraphs:
        if "[" in paragraph.text or "]" in paragraph.text:
            _convert_paragraph(paragraph, counters, occurrences)


def _iter_cell_paragraphs(table):
    """Cell paragraphs in reading order: rows top to bottom, cells left to right.

    A merged cell appears once per grid slot, so its paragraphs can be yielded more than once.
    That is safe -- including for the occurrence counters -- because conversion strips every
    bracket on the first visit and `_convert_paragraphs` skips bracket-free text.
    """
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_cell_paragraphs(nested)


_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

REFERENCE_TABLE_HEADER = "DATA TAG"
REFERENCE_HEADING = "Data Tag Reference"


def _drop_reference_table(doc) -> bool:
    """Delete the draft's Data Tag Reference table and its heading paragraph.

    The table is the doc team's authoring aid (every tag, its source, where it appears), not
    document content -- team decision 2026-08-31. Must run BEFORE conversion: with it gone,
    occurrence numbering starts at the first REAL use of a tag, which is what the
    FEASIBILITY_OCCURRENCES rules assume.
    """
    dropped = False
    for table in list(doc.tables):
        if table.rows and REFERENCE_TABLE_HEADER in table.rows[0].cells[0].text.upper():
            table._tbl.getparent().remove(table._tbl)
            dropped = True
    for paragraph in list(doc.paragraphs):
        if REFERENCE_HEADING.lower() in paragraph.text.lower():
            paragraph._element.getparent().remove(paragraph._element)
    return dropped


def _control_row(template_row):
    """A copy of `template_row` with every run removed, ready to hold one control tag."""
    tr = deepcopy(template_row._tr)
    for p in tr.findall(f".//{_NS}p"):
        for r in p.findall(f"{_NS}r"):
            p.remove(r)
    return tr


def _set_row_cell_texts(row, mapping: dict[str, str], literal: dict[str, str]) -> None:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            text = paragraph.text
            stripped = text.strip()
            if stripped in literal:
                _rewrite_paragraph(paragraph, literal[stripped])
                continue
            if REPEAT_NOTE in text:
                _rewrite_paragraph(paragraph, "")
                continue
            match = BRACKET.search(text)
            if match:
                tag = match.group(1)
                for key, expr in mapping.items():
                    if tag.endswith(key):
                        _rewrite_paragraph(paragraph, expr)
                        break


def _convert_loop_rows(table) -> bool:
    """Wrap this table's marker row in a {%tr%} loop, if it has one. True when converted."""
    converted = False
    for spec in LOOP_ROWS:
        for row in table.rows:
            if spec["marker"] in row.cells[0].text or any(
                    spec["marker"] in cell.text for cell in row.cells):
                _set_row_cell_texts(row, spec["cells"], spec.get("literal", {}))

                open_tr = _control_row(row)
                row._tr.addprevious(open_tr)
                _first_paragraph_text(open_tr, "{%%tr for r in %s %%}" % spec["var"])

                close_tr = _control_row(row)
                row._tr.addnext(close_tr)
                _first_paragraph_text(close_tr, "{%tr endfor %}")

                converted = True
                break
    return converted


def _first_paragraph_text(tr, text: str) -> None:
    p = tr.find(f".//{_NS}p")
    r = p.makeelement(f"{_NS}r", {})
    t = p.makeelement(f"{_NS}t", {})
    t.text = text
    r.append(t)
    p.append(r)


# The Monitoring Plan's ecosystem tables repeat identical tags once per ecosystem block, so
# each occurrence is routed by position into the `ecos` context list (3 entries, pathway card
# order: Forest, Mangrove, Peatland). Occurrence numbers are DOCUMENT ORDER, template-wide:
# occurrence 1 of each tag sits in the "Project Activity Type & Land Area" paragraph and stays
# generic; the tables follow. Re-derive these rules if the doc team restructures the template.
# Expressions must stay BRACKET-FREE (`eco0.label`, never `ecos[0].label`): a merged table cell
# is visited once per grid slot, and a revisit re-converts any text still containing `[`, which
# would wrap the index of a subscripted expression in a nested tag.
def _eco_field(field):
    #  n == 1: the prose paragraph -> generic. n in 2..7: two tables x 3 blocks -> (n-2) % 3.
    return lambda n: None if n == 1 else "eco%d.%s" % ((n - 2) % 3, field)


def _chosen_activities(n):
    # n == 1: prose. n in 2..10: Selected Activities table, 3 blocks x Protect/Manage/Restore.
    # n in 11..13: Monitoring Indicators, one per ecosystem, all interventions merged.
    if n == 1:
        return None
    if n <= 10:
        intervention = ("protect", "manage", "restore")[(n - 2) % 3]
        return "eco%d.activities_%s" % ((n - 2) // 3, intervention)
    return "eco%d.activities_all" % (n - 11)


# The feasibility template repeats `[Potential Benefit: X tonnes]` in avoided/sequestered pairs
# ("X tonnes avoided ... X tonnes sequestered", twice: Net Carbon Removal Estimates and Technical
# Feasibility). The Data Tag Reference table -- whose row used to be occurrence 1 -- is dropped
# before conversion, so numbering starts at the first real pair: odd occurrences are the avoided
# figure and even ones the sequestered figure.
FEASIBILITY_OCCURRENCES = {
    "Potential Benefit: X tonnes": lambda n: (
        "benefit_avoided_tco2e" if n % 2 == 1 else "benefit_sequestered_tco2e"),
    # "x Nature sub-components, x People sub-components and x Climate sub-components scored".
    "Potential Benefit: x": lambda n: (
        ("benefit_nature_count", "benefit_people_count",
         "benefit_climate_count")[(n - 1) % 3]),
}

MONITORING_OCCURRENCES = {
    "Threat: Forest / Mangrove / Peatland": _eco_field("label"),
    "NbS Pathway: hectare area eligible to protect": _eco_field("protect_ha"),
    "NbS Pathway: hectare area eligible to manage": _eco_field("manage_ha"),
    "NbS Pathway: hectare area eligible to restore": _eco_field("restore_ha"),
    "NbS Pathway: Chosen NbS Activities": _chosen_activities,
}

# The "Monitoring Indicators, Methods and Frequency" tables are header-only in the draft -- no
# data row and no tags -- so a loop row is INJECTED per table. The three tables appear in
# ecosystem order (Dryland, Mangrove, Peatland), matching eco0..eco2. Column order mirrors the
# draft's header: Benefit Category | Benefit | Indicators | Unit | frequency | Methodology.
INDICATOR_HEADER = "Benefit Category"
INDICATOR_COLUMNS = ("category", "benefit", "indicator", "unit", "freq", "source")


def _inject_indicator_rows(doc) -> int:
    """Append `{%tr for i in ecoN.indicator_rows %}` + data row + endfor to each indicator
    table, in document order. Returns how many tables were wired."""
    count = 0
    for table in doc.tables:
        # The header row that names the columns is the one containing INDICATOR_HEADER (it may
        # sit below a title row); clone it as the shape template so the injected row inherits
        # the table's grid and formatting.
        template_row = next((r for r in table.rows
                             if any(INDICATOR_HEADER in c.text for c in r.cells)), None)
        if template_row is None:
            continue

        data_tr = _control_row(template_row)
        template_row._tr.getparent().append(data_tr)
        for cell_tc, field in zip(data_tr.findall(f"{_NS}tc"), INDICATOR_COLUMNS):
            p = cell_tc.find(f"{_NS}p")
            r = p.makeelement(f"{_NS}r", {})
            t = p.makeelement(f"{_NS}t", {})
            t.text = "{{ i.%s }}" % field
            r.append(t)
            p.append(r)

        open_tr = _control_row(template_row)
        data_tr.addprevious(open_tr)
        _first_paragraph_text(open_tr, "{%%tr for i in eco%d.indicator_rows %%}" % count)

        close_tr = _control_row(template_row)
        data_tr.addnext(close_tr)
        _first_paragraph_text(close_tr, "{%tr endfor %}")
        count += 1
    return count


def convert(source_path: str, output_path: str, occurrences: dict | None = None) -> None:
    doc = Document(source_path)
    counters: dict = {}

    # Before anything counts occurrences: the authoring aid is not document content.
    _drop_reference_table(doc)

    # Body content in DOCUMENT ORDER -- occurrence counting depends on it. Loop rows are wrapped
    # before a table's paragraphs are converted, so the loop cells get their own expressions.
    for block in doc.iter_inner_content():
        if hasattr(block, "rows"):          # a Table
            _convert_loop_rows(block)
            _convert_paragraphs(_iter_cell_paragraphs(block), counters, occurrences)
        else:                               # a Paragraph
            _convert_paragraphs([block], counters, occurrences)

    for section in doc.sections:
        _convert_paragraphs(section.header.paragraphs, counters, occurrences)
        _convert_paragraphs(section.footer.paragraphs, counters, occurrences)

    if occurrences is MONITORING_OCCURRENCES:
        _inject_indicator_rows(doc)

    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "monitoring":
        convert("assets/Monitoring Plan_Template Draft.docx",
                "assets/monitoring_v3_template.docx", MONITORING_OCCURRENCES)
        print("converted: monitoring -> assets/monitoring_v3_template.docx")
    else:
        src = sys.argv[1] if len(sys.argv) > 1 else "assets/Feasibility Study_Template Draft.docx"
        out = sys.argv[2] if len(sys.argv) > 2 else "assets/feasibility_v3_template.docx"
        convert(src, out, FEASIBILITY_OCCURRENCES)
        print(f"converted: {src} -> {out}")
